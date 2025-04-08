from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes, ConversationHandler
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
import json
import sys
import time
import os
import asyncio

# Состояния разговора
NAME, BIRTHDATE, GOAL = range(3)

# Клавиатура разработчика
developer_keyboard = InlineKeyboardMarkup(
    inline_keyboard=[[
        InlineKeyboardButton(
            text="🛠️ GitHub",
            url="https://github.com/bayanist"
        ),
        InlineKeyboardButton(
            text="💬 Telegram",
            url="https://t.me/bayanman"
        )
    ]]
)

# Конфигурация API
TELEGRAM_TOKEN = "ТВОЙ"
GEMINI_API_KEY = "ТВОЙ"
GEMINI_API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Начало разговора и запрос ФИО"""
    # Отправляем сообщение о разработчике
    await update.message.reply_text(
        "💬 Разработчик: Титов Антон",
        reply_markup=developer_keyboard
    )
    
    # Отправляем основное сообщение
    await update.message.reply_text(
        "👋 Здравствуйте! Я помогу составить индивидуальный план работы.\n"
        "Для начала, пожалуйста, введите Ф.И.О. ребёнка:"
    )
    return NAME

async def get_name(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Сохранение ФИО и запрос даты рождения"""
    context.user_data['name'] = update.message.text
    await update.message.reply_text("Введите дату рождения (дд.мм.гггг):")
    return BIRTHDATE

async def get_birthdate(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Сохранение даты рождения и запрос цели"""
    context.user_data['birthdate'] = update.message.text
    await update.message.reply_text(
        "Введите цель работы:\n"
        "Например: Повышение оценки через создание позитивного 'Я' образа"
    )
    return GOAL

async def generate_content_with_progress(update, prompt):
    """Генерация контента с отображением прогресса"""
    try:
        status_message = await update.message.reply_text("⏳ Подключаемся к API...")
        
        # Подготовка данных для запроса
        payload = {
            "contents": [{
                "parts":[{
                    "text": prompt
                }]
            }]
        }
        
        headers = {
            'Content-Type': 'application/json'
        }

        # Тестовый запрос
        try:
            test_payload = {
                "contents": [{
                    "parts":[{
                        "text": "test connection"
                    }]
                }]
            }
            
            test_response = requests.post(
                GEMINI_API_URL,
                headers=headers,
                json=test_payload,
                timeout=10
            )
            test_response.raise_for_status()
            await status_message.edit_text("✅ Подключение установлено\n⏳ Генерируем план...")
        except Exception as e:
            await status_message.edit_text(f"❌ Ошибка подключения: {str(e)}")
            print(f"Детали ошибки подключения: {test_response.text if 'test_response' in locals() else 'Нет ответа'}")
            return None

        # Основной запрос
        try:
            response = requests.post(
                GEMINI_API_URL,
                headers=headers,
                json=payload,
                timeout=30
            )
            response.raise_for_status()
            
            response_data = response.json()
            print(f"Ответ API: {json.dumps(response_data, indent=2, ensure_ascii=False)}")
            
            if 'candidates' in response_data and len(response_data['candidates']) > 0:
                content = response_data['candidates'][0].get('content', {})
                parts = content.get('parts', [])
                if parts and 'text' in parts[0]:
                    text = parts[0]['text']
                    await status_message.edit_text("✅ План успешно сгенерирован!")
                    return text
            
            raise Exception("Неверный формат ответа от API")
            
        except Exception as e:
            await status_message.edit_text(f"❌ Ошибка генерации: {str(e)}")
            print(f"Детали ошибки генерации: {response.text if 'response' in locals() else 'Нет ответа'}")
            return None

    except Exception as e:
        await update.message.reply_text(f"❌ Произошла ошибка: {str(e)}")
        return None

async def generate_plan(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Генерация плана с помощью Gemini API"""
    try:
        # Если цель уже есть в данных (при reset) - используем её, иначе берём из сообщения
        if 'goal' in context.user_data and update.message.text == '/reset':
            goal = context.user_data['goal']
        else:
            goal = update.message.text
            context.user_data['goal'] = goal
        
        prompt = f"""
        На основе цели '{goal}', составь таблицу с тремя колонками:
        1. Направление (какие аспекты работы подходят под эту цель).
        2. Коррекционные задачи (что именно корректируем/развиваем).
        3. Темы занятий (конкретные темы, соответствующие задачам).

        Пример формата ответа:
        ---
        Коррекция и развитие эмоционально-волевой сферы | Развитие самоконтроля, снижение тревожности | Как справляться со стрессом?, Что такое сила воли?
        ---
        Ответ дай только в таком формате, без заголовков.
        """

        generated_text = await generate_content_with_progress(update, prompt)
        
        if not generated_text:
            await update.message.reply_text("❌ Не удалось сгенерировать план. Попробуйте позже или используйте /start для нового запроса.")
            return ConversationHandler.END

        # Создаем документ
        await update.message.reply_text("📝 Создаём документ...")
        doc = Document()
        
        # Добавляем шапку как в примере
        school_name = doc.add_paragraph('ГБОУ «Реабилитационная школа-интернат «Восхождение»»')
        school_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        school_name.runs[0].bold = True
        
        title = doc.add_paragraph('Индивидуальный план работы педагога-психолога')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title.runs[0].bold = True
        
        teacher_name = doc.add_paragraph('Тевфиковой Анны Николаевны на 2024-2025 учебный год')
        teacher_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        teacher_name.runs[0].bold = True
        
        # Добавляем информацию о ребенке
        doc.add_paragraph(f'Ф.И.О.: {context.user_data["name"]}')
        doc.add_paragraph(f'Класс: [укажите класс]')  # Можно добавить запрос класса в диалоге
        doc.add_paragraph(f'Дата рождения: {context.user_data["birthdate"]}')
        doc.add_paragraph(f'Цель: {context.user_data["goal"]}')

        # Добавляем таблицу
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        # Заголовки таблицы
        headers = ['Направление', 'Коррекционные задачи', 'Темы', 'Кол-во занятий', 'Динамика (январь)', 'Динамика (май)']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        # Заполняем таблицу
        lines = generated_text.split("\n")
        for line in lines:
            if line.strip():
                parts = line.split("|")
                if len(parts) == 3:
                    row_cells = table.add_row().cells
                    row_cells[0].text = parts[0].strip()
                    row_cells[1].text = parts[1].strip()
                    row_cells[2].text = parts[2].strip()
                    row_cells[3].text = " "
                    row_cells[4].text = " "
                    row_cells[5].text = " "

        # Сохраняем документ
        filename = f'Индивидуальный_план_{context.user_data["name"].replace(" ", "_")}.docx'
        doc.save(filename)

        # Отправляем файл
        await update.message.reply_text("📤 Отправляю файл...")
        await update.message.reply_document(
            document=open(filename, 'rb'),
            filename=filename
        )
        
        # Удаляем временный файл
        os.remove(filename)
        
        await update.message.reply_text(
            "✅ Готово! Можете начать новый план командой /start\n🔄 Пересоздать план /reset"
        )
        
        return ConversationHandler.END

    except Exception as e:
        await update.message.reply_text(f"❌ Произошла ошибка: {str(e)}\nПопробуйте начать сначала с помощью команды /start")
        return ConversationHandler.END

async def reset(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Пересоздание плана с использованием сохраненных данных"""
    if not all(key in context.user_data for key in ['name', 'birthdate', 'goal']):
        await update.message.reply_text(
            "❌ Нет сохраненных данных. Пожалуйста, начните сначала с помощью команды /start"
        )
        return ConversationHandler.END

    await update.message.reply_text("🔄 Пересоздаю план с сохраненными данными...")
    return await generate_plan(update, context)

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отмена операции"""
    await update.message.reply_text("❌ Операция отменена. Для начала нового плана используйте /start")
    return ConversationHandler.END

def main():
    """Запуск бота"""
    try:
        print("🤖 Запуск бота...")
        print("📡 Проверка подключения к Gemini API...")
        print(f"URL API: {GEMINI_API_URL}")
        
        # Тестовое подключение к API
        try:
            test_payload = {
                "contents": [{
                    "parts":[{
                        "text": "test connection"
                    }]
                }]
            }
            
            test_response = requests.post(
                GEMINI_API_URL,
                headers={'Content-Type': 'application/json'},
                json=test_payload,
                timeout=10
            )
            test_response.raise_for_status()
            response_data = test_response.json()
            print(f"✅ Подключение к Gemini API успешно!")
            print(f"Тестовый ответ: {json.dumps(response_data, indent=2, ensure_ascii=False)}")
        except Exception as e:
            print(f"❌ Ошибка подключения к Gemini API: {str(e)}")
            print(f"Детали ошибки: {test_response.text if 'test_response' in locals() else 'Нет ответа'}")
            sys.exit(1)

        # Создаем приложение
        application = Application.builder().token(TELEGRAM_TOKEN).build()

        # Добавляем обработчик разговора
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', start), CommandHandler('reset', reset)],
            states={
                NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_name)],
                BIRTHDATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_birthdate)],
                GOAL: [MessageHandler(filters.TEXT & ~filters.COMMAND, generate_plan)],
            },
            fallbacks=[CommandHandler('cancel', cancel)],
        )

        application.add_handler(conv_handler)

        # Запускаем бота
        print("✅ Бот готов к работе!")
        application.run_polling(allowed_updates=Update.ALL_TYPES)

    except Exception as e:
        print(f"❌ Ошибка при запуске бота: {str(e)}")

if __name__ == '__main__':
    main()
